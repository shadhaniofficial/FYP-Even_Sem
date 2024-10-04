from flask import Flask, request, render_template, redirect, url_for
import os
import sqlite3
from gtts import gTTS
import tempfile
from docx import Document
from googletrans import Translator

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOCX_FILE_PATH'] = 'uploads/fess101.docx'  # Path to your uploaded file

# Ensure the upload folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def create_database(db_name="my_paragraphs_db.db"):
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS paragraphs_table11 (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            page_no INTEGER,
            paragraph_no INTEGER,
            paragraph TEXT,
            audio_path TEXT,
            chapter_name TEXT,
            chapter_no INTEGER
        )
    ''')
    conn.commit()
    conn.close()

def create_table_if_not_exists(db_name):
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS paragraphs_table111 (
            page_no INTEGER,
            paragraph_no INTEGER,
            paragraph TEXT,
            audio_path TEXT,
            chapter_name TEXT,
            chapter_no INTEGER
                   
        );
    ''')
    conn.commit()
    conn.close()

# Call this function at the start of your application
create_table_if_not_exists("my_paragraphs_db.db")


def store_paragraph_and_audio_in_db(page_no, paragraph_no, paragraph, audio_path, chapter_name, chapter_no, db_name="my_paragraphs_db.db"):
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO paragraphs_table11 (page_no, paragraph_no, paragraph, audio_path, chapter_name, chapter_no)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (page_no, paragraph_no, paragraph, audio_path, chapter_name, chapter_no))
    conn.commit()
    conn.close()


def convert_text_to_audio(text, language='ta'):
    """
    Converts text to speech using gTTS and saves the audio file in the static/audios folder.
    
    Args:
        text: The text to be converted into speech.
        language: Language of the text (default is English).

    Returns:
        The filename of the saved audio file.
    """
    tts = gTTS(text=text, lang=language, slow=False)
    
    # Use first 10 characters of text for the filename, ensuring it's safe
    safe_text = ''.join(e for e in text[:10] if e.isalnum())  # Clean text for filename
    audio_filename = f"{safe_text}.mp3"  
    audio_path = os.path.join('static', 'audios', audio_filename)

    # Save the audio file in static/audios
    tts.save(audio_path)


    print("hh : " + audio_path)

    return audio_path  # Return the full path to the audio file
    
@app.route('/word-count', methods=['POST'])
def count_word_occurrences_route():
    """
    Count occurrences of a word in the paragraphs stored in the database.
    """
    word_to_find = request.form.get('word_to_find')  # Fetch word from form input
    db_name = "my_paragraphs_db.db"  # Name of the database

    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()

    # Fetch all paragraphs from the database
    cursor.execute('SELECT paragraph FROM paragraphs_table11')
    paragraphs = cursor.fetchall()
    conn.close()

    # Count occurrences of the word in all paragraphs
    word_count = 0
    for paragraph_tuple in paragraphs:
        paragraph = paragraph_tuple[0]  # Extract the paragraph from the tuple
        word_count += paragraph.lower().count(word_to_find.lower())  # Case-insensitive count

    # Pass word and count to template for rendering
    return render_template('word_count.html', word=word_to_find, count=word_count)


def fetch_paragraph(page_no, paragraph_no, db_name="my_paragraphs_db.db"):
    """
    Fetches the paragraph and associated audio path from the database.
    If the audio is not present, it will be generated and saved in static/audios.
    """
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()

    # Fetch the paragraph and audio path from the database
    cursor.execute('SELECT paragraph, audio_path FROM paragraphs_table11 WHERE page_no = ? AND paragraph_no = ?', (page_no, paragraph_no))
    row = cursor.fetchone()
    conn.close()

    if row:
        paragraph, audio_path = row

        # Check if audio exists or not in the static/audios folder
        if not os.path.exists(audio_path):
            # Audio doesn't exist, generate it again
            audio_path = convert_text_to_audio(paragraph)
            
            # Optionally, you can update the database with the new audio path
            conn = sqlite3.connect(db_name)
            cursor = conn.cursor()
            cursor.execute('UPDATE paragraphs_table11 SET audio_path = ? WHERE page_no = ? AND paragraph_no = ?', 
                           (audio_path, page_no, paragraph_no))
            conn.commit()
            conn.close()

        return paragraph, audio_path  # Return both paragraph and audio path
    else:
        return None, None  # Return None if no row is found
def fetch_paragraphs_by_keyword(keyword, db_name="my_paragraphs_db.db"):
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    cursor.execute('SELECT paragraph, audio_path FROM paragraphs_table11 WHERE paragraph LIKE ?', ('%' + keyword + '%',))
    rows = cursor.fetchall()
    conn.close()
    return rows

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/retrieve', methods=['POST'])
def retrieve_paragraph():
    page_no = request.form.get('page_no')
    paragraph_no = request.form.get('paragraph_no')
    paragraph_data = fetch_paragraph(page_no, paragraph_no)

    if paragraph_data:
        paragraph, audio_path = paragraph_data
        return render_template('paragraph.html', paragraph=paragraph, audio_path=audio_path)
    else:
        return "Paragraph not found."

@app.route('/search', methods=['POST'])
def search_paragraphs():
    keyword = request.form.get('keyword')
    results = fetch_paragraphs_by_keyword(keyword)

    # Process results and check if audio file exists
    for i, (text, audio_path) in enumerate(results):
        if not audio_path or not os.path.exists(audio_path):
            # If audio file doesn't exist, generate it
            new_audio_path = convert_text_to_audio(text, "en")
            results[i] = (text, new_audio_path)  # Update with new audio path

    if results:
        return render_template('search_results.html', results=results)
    else:
        return "No paragraphs found containing the keyword."

def extract_text_from_word(docx_file_path, chapter_name, chapter_no, db_name="my_paragraphs_db.db"):
    """
    Extracts text paragraph-wise from a Word file (.docx), simulates page numbers, and stores it in the database
    along with chapter information.

    Args:
        docx_file_path: The path to the Word file.
        chapter_name: The name of the chapter.
        chapter_no: The chapter number.
        db_name: The name of the SQLite database file.

    Returns:
        A string containing all the extracted text.
    """
    # Connect to the database and get the highest page number
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    cursor.execute("SELECT MAX(page_no) FROM paragraphs_table11")
    result = cursor.fetchone()
    max_page_no = result[0] if result[0] else 0  # If no pages, start at page 1
    conn.close()

    # Start page numbering from the max page number + 1
    page_no = max_page_no 
    paragraph_counter = 0
    full_text = []

    # Open the Word document and extract the paragraphs
    doc = Document(docx_file_path)

    # Collect text from the document including '$$'
    for para in doc.paragraphs:
        paragraph_text = para.text.strip()

        if paragraph_text:  # If the paragraph is not empty
            full_text.append(paragraph_text)

    # Combine all extracted text
    combined_text = "\n".join(full_text)

    # Split by '$$' for page breaks and by '$' for paragraphs
    pages = combined_text.split('$$')

    for page in pages:
        paragraphs = page.split('$')

        for sub_paragraph in paragraphs:
            clean_paragraph = sub_paragraph.strip()

            if clean_paragraph:  # If the sub-paragraph is not empty
                audio_path = convert_text_to_audio(clean_paragraph)

                # Store the paragraph with page number, paragraph number, and chapter info
                store_paragraph_and_audio_in_db(page_no, paragraph_counter + 1, clean_paragraph, audio_path, chapter_name, chapter_no, db_name)

                # Update the paragraph counter
                paragraph_counter += 1

        # Move to the next page after each '$$'
        page_no += 1
        paragraph_counter = 0  # Reset the paragraph counter for the new page

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        file = request.files['file']
        chapter_name = request.form['chapter_name']
        chapter_no = request.form['chapter_no']

        if file and file.filename.endswith('.docx'):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)

            # Extract text from the uploaded Word file and store it in the database
            extract_text_from_word(file_path, chapter_name, chapter_no)

            return redirect(url_for('index'))

    return render_template('upload.html')
@app.route('/fetch_data')
def fetch_data():
    conn = sqlite3.connect("my_paragraphs_db.db")
    cursor = conn.cursor()
    cursor.execute('SELECT page_no,paragraph_no,paragraph,audio_path,chapter_name,chapter_no FROM paragraphs_table11')
    rows = cursor.fetchall()
    conn.close()
    return render_template('data.html', rows=rows)
    




def get_db_connection(db_name="my_paragraphs_db.db"):
    conn = sqlite3.connect(db_name)
    conn.row_factory = sqlite3.Row  # To access columns by name
    return conn
def extract_topics_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    topics = {}
    current_topic = None
    for idx, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):  # Identify topic headings
            current_topic = para.text.strip()
            topics[current_topic] = idx  # Store the paragraph index of the heading
    return topics

# Translate text to Tamil
def translate_text_to_tamil(text):
    translator = Translator()
    return translator.translate(text, dest='ta').text
    
@app.route('/trans', methods=['GET', 'POST'])
def upload_file1():
    if request.method == 'POST':
        # file = request.files['file']
        # if file and file.filename.endswith('.docx'):

        file_names = [];
        # Append all the file names in upload_folder to the file names array
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            if filename.endswith('.docx'):
                file_names.append(filename)

        print(file_names)

        mapped_files = {}
        
        for f_name in file_names:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], f_name)
            # file.save(file_path)
            
            # Extract topics from the DOCX file
            topics = extract_topics_from_docx(file_path)

            # topics = [topic.replace('$', '') for topic in topics]            

            mapped_files[f_name] = topics

            # Store topics in the database (optional)

        return render_template('select_topics.html', topics=mapped_files)
    
    return render_template('upload.html')

@app.route('/select-topics', methods=['POST'])
def translate_topics():
    selected_topics = request.form.getlist('topics')
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], request.form['file_name'])  # Change to the uploaded file

    # Extract topics and paragraphs from the DOCX file
    topics = extract_topics_from_docx(file_path)
    doc = Document(file_path)
    translated_content = {}
    
    # Fetch and translate content under selected topics
    for selected_topic in selected_topics:
        if selected_topic in topics:
            translated_content[selected_topic] = []
            current_topic = selected_topic
            start_index = topics[current_topic]

            # Find paragraphs under the current topic
            for idx in range(start_index + 1, len(doc.paragraphs)):
                para = doc.paragraphs[idx]
                if para.style.name.startswith('Heading'):
                    break  # Stop if we reach the next topic

                if para.text.strip():  # Only translate non-empty paragraphs
                    translated_paragraph = translate_text_to_tamil(para.text.strip())
                    print(translated_paragraph)
                    translated_content[selected_topic].append(translated_paragraph)
                    audio_path = convert_text_to_audio(translated_paragraph)  # Save audio for each translated paragraph
                    break;


        return render_template('translated_topics.html', content=translated_content, audio_path=audio_path)

def truncate_table(db_name="my_paragraphs_db.db", table_name="paragraphs_table11"):
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    cursor.execute(f'DELETE FROM {table_name}')
    conn.commit()
    cursor.execute(f'VACUUM')  # Optional: to reclaim space in SQLite after deletion
    conn.close()

@app.route('/truncate', methods=['POST'])
def truncate_paragraphs_table():
    truncate_table()
    return "Table truncated successfully."


if __name__ == '__main__':
    app.run(debug=True)
